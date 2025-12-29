#!/usr/bin/env python3
"""
ZoteroResearcher Common Module

Shared base class and utilities for all ZoteroResearcher workflows.
"""

import requests
import trafilatura
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
from typing import Optional, Dict, List, Tuple, Any
from anthropic import Anthropic

# Handle both relative and absolute imports
try:
    from .zotero_base import ZoteroBaseProcessor
    from .zr_llm_client import ZRLLMClient
except ImportError:
    from zotero_base import ZoteroBaseProcessor
    from zr_llm_client import ZRLLMClient


def validate_project_name(name: str) -> str:
    """
    Validate and sanitize project name.

    Args:
        name: Raw project name from user

    Returns:
        Validated project name

    Raises:
        ValueError: If project name is invalid
    """
    if not name:
        raise ValueError("Project name cannot be empty")

    name = name.strip()

    if not name:
        raise ValueError("Project name cannot be empty or whitespace only")

    if len(name) > 50:
        raise ValueError(f"Project name too long: '{name}' (max 50 characters)")

    # Check for problematic characters
    # Zotero handles most characters fine, but let's be cautious with special chars
    problematic_chars = ['【', '】', '\n', '\r', '\t']
    for char in problematic_chars:
        if char in name:
            raise ValueError(f"Project name contains invalid character: '{char}'")

    return name


class ZoteroResearcherBase(ZoteroBaseProcessor):
    """Base class for ZoteroResearcher workflows with shared functionality."""

    # Content truncation limits for LLM processing
    GENERAL_SUMMARY_CHAR_LIMIT = 500000   # Phase 1: General summaries
    TARGETED_SUMMARY_CHAR_LIMIT = 500000  # Phase 2: Targeted summaries

    def __init__(
        self,
        library_id: str,
        library_type: str,
        api_key: str,
        anthropic_api_key: str,
        project_name: str = None,
        force_rebuild: bool = False,
        verbose: bool = False,
        use_cache: bool = False,
        cache_manager = None
    ):
        """
        Initialize the Zotero researcher base.

        Args:
            library_id: Your Zotero user ID or group ID
            library_type: 'user' or 'group'
            api_key: Your Zotero API key
            anthropic_api_key: Anthropic API key for Claude
            project_name: Name of the research project (used for organizing subcollections and notes)
            force_rebuild: If True, force rebuild of existing general summaries (default: False)
            verbose: If True, show detailed information about all child items
            use_cache: If True, use local cache instead of API calls (default: False)
            cache_manager: ZoteroCacheManager instance (required if use_cache=True)
        """
        # Initialize base class
        super().__init__(library_id, library_type, api_key, verbose)

        # Validate and store project name
        self.project_name = validate_project_name(project_name) if project_name else None

        # Cache configuration
        self.use_cache = use_cache
        self.cache_manager = cache_manager

        # Researcher-specific configuration
        self.anthropic_client = Anthropic(api_key=anthropic_api_key)

        # Operational flags
        self.force_rebuild = force_rebuild

        # Content loaded from Zotero (populated during operations)
        self.research_brief = ""
        self.project_overview = ""
        self.tags = []

        # Initialize configurable parameters with defaults
        # (These can be overridden by project config file)
        self.max_workers = 20
        self.rate_limit_delay = 0.1
        self.general_summary_char_limit = self.GENERAL_SUMMARY_CHAR_LIMIT
        self.targeted_summary_char_limit = self.TARGETED_SUMMARY_CHAR_LIMIT
        self.relevance_threshold = 6
        self.max_sources = 50
        self.use_sonnet = False
        self.synthesis_enabled = True  # Default: enabled

        # LLM model configuration
        self.haiku_model = "claude-haiku-4-5-20251001"
        self.sonnet_model = "claude-sonnet-4-5-20250929"

        # Default to Haiku for detailed summaries (cost-efficient)
        # Use Sonnet only when use_sonnet=True in config (production mode)
        self.summary_model = self.haiku_model

        # Initialize centralized LLM client for all API calls
        self.llm_client = ZRLLMClient(
            anthropic_client=self.anthropic_client,
            default_model=self.haiku_model,
            verbose=verbose
        )

    def _get_subcollection_name(self) -> str:
        """Get project-specific subcollection name."""
        if not self.project_name:
            raise ValueError("Project name is required but not set")
        return f"【ZResearcher: {self.project_name}】"

    def _get_project_overview_note_title(self) -> str:
        """Get project-specific overview note title."""
        if not self.project_name:
            raise ValueError("Project name is required but not set")
        return f"【Project Overview】"

    def _get_research_tags_note_title(self) -> str:
        """Get project-specific tags note title."""
        if not self.project_name:
            raise ValueError("Project name is required but not set")
        return f"【Research Tags】"

    def _get_research_brief_note_title(self) -> str:
        """Get project-specific brief note title."""
        if not self.project_name:
            raise ValueError("Project name is required but not set")
        return f"【Research Brief】"

    def _get_query_request_note_title(self) -> str:
        """Get project-specific query request note title (for File Search)."""
        if not self.project_name:
            raise ValueError("Project name is required but not set")
        return f"【Query Request】"

    def _get_project_config_note_title(self) -> str:
        """Get project-specific config note title."""
        if not self.project_name:
            raise ValueError("Project name is required but not set")
        return f"【Project Config】"

    def _get_summary_note_prefix(self) -> str:
        """Get project-specific summary note prefix (used as note title/heading)."""
        if not self.project_name:
            raise ValueError("Project name is required but not set")
        return f"【ZResearcher Summary: {self.project_name}】"

    def _get_default_config_template(self) -> str:
        """Get the default project configuration template."""
        return """# ZResearcher Project Configuration
# Edit values below to customize this project's behavior
# Lines starting with # are comments and will be ignored

# ============================================================
# Performance Settings
# ============================================================
max_workers=20
rate_limit_delay=0.1

# ============================================================
# Content Truncation Limits (characters)
# ============================================================
general_summary_char_limit=500000
targeted_summary_char_limit=500000

# ============================================================
# Relevance & Filtering
# ============================================================
relevance_threshold=6
max_sources=50

# ============================================================
# LLM Model Configuration
# ============================================================
use_sonnet=false
haiku_model=claude-haiku-4-5
sonnet_model=claude-sonnet-4-5

# ============================================================
# Research Synthesis
# ============================================================
generate_synthesis=true

# ============================================================
# Gemini File Search Configuration
# ============================================================
gemini_file_search_model=gemini-2.5-pro

# ============================================================
# Gemini File API State (managed automatically)
# DO NOT EDIT - these are updated by --file-search
# ============================================================
gemini_uploaded_files={}

# ============================================================
# Notes:
# - Boolean values: true/false (case insensitive)
# - Integer values: whole numbers
# - Float values: decimal numbers (use . not ,)
# - Changes take effect on next build/query operation
# - Restart not required - just re-run the command
# ============================================================
"""

    def extract_metadata(self, item: Dict) -> Dict:
        """
        Extract metadata from a Zotero item.

        Args:
            item: The Zotero item

        Returns:
            Dict with metadata fields (title, authors, date, publication, url, itemType)
        """
        item_data = item['data']

        # Extract basic fields
        metadata = {
            'title': item_data.get('title', 'Untitled'),
            'date': item_data.get('date', 'Unknown date'),
            'publication': item_data.get('publicationTitle', item_data.get('bookTitle', '')),
            'url': item_data.get('url', ''),
            'itemType': item_data.get('itemType', 'unknown')
        }

        # Extract authors/creators
        creators = item_data.get('creators', [])
        if creators:
            authors = []
            for creator in creators:
                if 'lastName' in creator:
                    if 'firstName' in creator:
                        authors.append(f"{creator['firstName']} {creator['lastName']}")
                    else:
                        authors.append(creator['lastName'])
                elif 'name' in creator:
                    authors.append(creator['name'])
            metadata['authors'] = ', '.join(authors) if authors else 'Unknown author'
        else:
            metadata['authors'] = 'Unknown author'

        return metadata

    def extract_text_from_html(self, html_content, attachment_url: Optional[str] = None) -> Optional[str]:
        """
        Extract text from HTML content using Trafilatura.
        Reused from summarize_sources.py.

        Args:
            html_content: HTML content as bytes or str
            attachment_url: Optional URL to fetch from if bytes fail

        Returns:
            Extracted text, or None if extraction fails
        """
        try:
            # Convert to string if bytes
            if isinstance(html_content, bytes):
                html_string = html_content.decode('utf-8', errors='ignore')
            else:
                html_string = html_content

            # Use Trafilatura for extraction
            markdown = trafilatura.extract(
                html_string,
                output_format='markdown',
                include_links=True,
                include_images=False,
                include_tables=True
            )

            if markdown:
                return markdown.strip()

            # If Trafilatura fails and we have a URL, try fetching directly
            if attachment_url and not markdown:
                print("  ⚠️  Trying to fetch from URL...")
                response = requests.get(attachment_url, timeout=30)
                response.raise_for_status()
                markdown = trafilatura.extract(
                    response.text,
                    output_format='markdown',
                    include_links=True,
                    include_images=False,
                    include_tables=True
                )
                if markdown:
                    return markdown.strip()

            return None

        except Exception as e:
            print(f"  ❌ Error extracting HTML: {e}")
            return None

    def extract_text_from_pdf(self, pdf_content: bytes) -> Optional[str]:
        """
        Extract text from a PDF using PyMuPDF.
        Reused from summarize_sources.py.

        Args:
            pdf_content: The PDF file content as bytes

        Returns:
            Extracted text as string, or None if extraction failed
        """
        try:
            # Open PDF from bytes
            pdf_document = fitz.open(stream=pdf_content, filetype="pdf")

            extracted_text = []
            total_pages = len(pdf_document)

            # Extract text from each page
            for page_num in range(total_pages):
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    extracted_text.append(text)

            pdf_document.close()

            full_text = "\n\n".join(extracted_text) if extracted_text else None

            # Check if PDF is likely scanned (very little text)
            if full_text and total_pages > 0:
                avg_chars_per_page = len(full_text) / total_pages
                if avg_chars_per_page < 100:
                    print(f"  ⚠️  Warning: PDF appears to be scanned (low text density)")

            return full_text

        except Exception as e:
            print(f"  ❌ Error extracting PDF text: {e}")
            return None

    def extract_text_from_txt(self, txt_content: bytes) -> Optional[str]:
        """
        Extract text from a plain text file.

        Args:
            txt_content: The text file content as bytes

        Returns:
            Extracted text as string, or None if extraction failed
        """
        try:
            # Try UTF-8 first, then fall back to other encodings
            try:
                text = txt_content.decode('utf-8')
            except UnicodeDecodeError:
                # Try common alternative encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = txt_content.decode(encoding)
                        print(f"  ℹ️  Decoded using {encoding} encoding")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # If all else fails, decode with errors='replace'
                    text = txt_content.decode('utf-8', errors='replace')
                    print(f"  ⚠️  Warning: Some characters could not be decoded")

            return text.strip() if text else None

        except Exception as e:
            print(f"  ❌ Error extracting text: {e}")
            return None

    def extract_text_from_docx(self, docx_content: bytes) -> Optional[str]:
        """
        Extract text from a DOCX file with structure preservation.
        Converts headings, tables, and lists to markdown-like format.

        Args:
            docx_content: The DOCX file content as bytes

        Returns:
            Extracted text as markdown-like string, or None if extraction failed
        """
        try:
            # Open DOCX from bytes using BytesIO
            doc = Document(BytesIO(docx_content))

            extracted_parts = []

            # Extract text from paragraphs with style preservation
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue

                # Detect heading styles and convert to markdown
                style_name = paragraph.style.name.lower()
                if 'heading 1' in style_name:
                    extracted_parts.append(f"\n# {text}\n")
                elif 'heading 2' in style_name:
                    extracted_parts.append(f"\n## {text}\n")
                elif 'heading 3' in style_name:
                    extracted_parts.append(f"\n### {text}\n")
                elif 'heading 4' in style_name:
                    extracted_parts.append(f"\n#### {text}\n")
                elif 'heading 5' in style_name:
                    extracted_parts.append(f"\n##### {text}\n")
                elif 'heading 6' in style_name:
                    extracted_parts.append(f"\n###### {text}\n")
                else:
                    extracted_parts.append(text)

            # Extract text from tables
            if doc.tables:
                for table in doc.tables:
                    extracted_parts.append("\n")  # Add spacing before table
                    for i, row in enumerate(table.rows):
                        row_text = []
                        for cell in row.cells:
                            cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                            row_text.append(cell_text)

                        # Create markdown-style table
                        extracted_parts.append("| " + " | ".join(row_text) + " |")

                        # Add separator after header row
                        if i == 0:
                            extracted_parts.append("| " + " | ".join(["---"] * len(row_text)) + " |")

                    extracted_parts.append("\n")  # Add spacing after table

            full_text = "\n".join(extracted_parts) if extracted_parts else None

            # Check if document has meaningful content
            if full_text and len(full_text.strip()) < 50:
                print(f"  ⚠️  Warning: DOCX appears to have very little text content")

            return full_text.strip() if full_text else None

        except Exception as e:
            # Check for specific error types
            error_msg = str(e).lower()
            if 'password' in error_msg or 'encrypted' in error_msg:
                print(f"  ❌ Error: DOCX file appears to be password-protected")
            elif 'not a zip file' in error_msg or 'bad zipfile' in error_msg:
                # This likely means it's a legacy .doc file, not .docx
                print(f"  ❌ Error: File appears to be legacy .doc format (not supported)")
                print(f"  ℹ️  Tip: Convert to .docx format for extraction support")
            else:
                print(f"  ❌ Error extracting DOCX text: {e}")
            return None

    def get_source_content(self, item: Dict) -> Tuple[Optional[str], Optional[str]]:
        """
        Get content from a source using priority order:
        0. Cached extracted content (if using cache)
        1. HTML snapshot (Trafilatura)
        2. PDF attachment (PyMuPDF)
        3. DOCX attachment (python-docx)
        4. TXT attachment (plain text)
        5. URL fetch (for webpage items)

        Args:
            item: The Zotero item

        Returns:
            Tuple of (content_text, content_type) or (None, None) if extraction fails
        """
        item_key = item['key']
        item_data = item['data']
        item_title = item_data.get('title', 'Untitled')

        # Priority 0: Use cached extracted content if available
        if self.use_cache and self.cache_manager:
            cached_content = self.cache_manager.get_cached_content(item_key)
            if cached_content:
                if self.verbose:
                    print(f"  📦 Using cached content ({len(cached_content):,} chars)")
                return cached_content, "CACHE"

        # Get attachments
        attachments = self.get_item_attachments(item_key)

        if attachments:
            # Priority 1: Try HTML attachment
            for attachment in attachments:
                if self.is_html_attachment(attachment):
                    attachment_title = attachment['data'].get('title', 'Untitled')
                    attachment_key = attachment['key']
                    attachment_url = attachment['data'].get('url')

                    print(f"  📄 Found HTML attachment: {attachment_title}")
                    print(f"  📥 Downloading and extracting...")

                    html_content = self.download_attachment(attachment_key)
                    if html_content:
                        extracted = self.extract_text_from_html(html_content, attachment_url)
                        if extracted:
                            return extracted, "HTML"

            # Priority 2: Try PDF attachment
            for attachment in attachments:
                if self.is_pdf_attachment(attachment):
                    attachment_title = attachment['data'].get('title', 'Untitled')
                    attachment_key = attachment['key']

                    print(f"  📄 Found PDF attachment: {attachment_title}")
                    print(f"  📥 Downloading and extracting...")

                    pdf_content = self.download_attachment(attachment_key)
                    if pdf_content:
                        extracted = self.extract_text_from_pdf(pdf_content)
                        if extracted:
                            return extracted, "PDF"

            # Priority 3: Try DOCX attachment
            for attachment in attachments:
                if self.is_docx_attachment(attachment):
                    attachment_title = attachment['data'].get('title', 'Untitled')
                    attachment_key = attachment['key']

                    print(f"  📄 Found DOCX attachment: {attachment_title}")
                    print(f"  📥 Downloading and extracting...")

                    docx_content = self.download_attachment(attachment_key)
                    if docx_content:
                        extracted = self.extract_text_from_docx(docx_content)
                        if extracted:
                            return extracted, "DOCX"

            # Priority 4: Try TXT attachment
            for attachment in attachments:
                if self.is_txt_attachment(attachment):
                    attachment_title = attachment['data'].get('title', 'Untitled')
                    attachment_key = attachment['key']

                    print(f"  📄 Found TXT attachment: {attachment_title}")
                    print(f"  📥 Downloading...")

                    txt_content = self.download_attachment(attachment_key)
                    if txt_content:
                        extracted = self.extract_text_from_txt(txt_content)
                        if extracted:
                            return extracted, "TXT"

        # Priority 5: Try fetching from URL (for webpage items)
        item_url = item_data.get('url')
        if item_url and item_data.get('itemType') == 'webpage':
            print(f"  🌐 Fetching from URL: {item_url}")
            try:
                response = requests.get(item_url, timeout=30)
                response.raise_for_status()
                markdown = trafilatura.extract(
                    response.text,
                    output_format='markdown',
                    include_links=True,
                    include_images=False,
                    include_tables=True
                )
                if markdown:
                    return markdown.strip(), "URL"
            except Exception as e:
                print(f"  ❌ Error fetching URL: {e}")

        return None, None

    def load_project_config_from_zotero(self, collection_key: str) -> Dict[str, Any]:
        """
        Load and parse project configuration from Zotero note.

        Args:
            collection_key: Parent collection key

        Returns:
            Dict with parsed configuration values (empty dict if not found)

        Raises:
            FileNotFoundError: If subcollection or config note not found
        """
        subcollection_name = self._get_subcollection_name()
        note_title = self._get_project_config_note_title()

        # Get project-specific subcollection
        subcollection_key = self.get_subcollection(collection_key, subcollection_name)
        if not subcollection_key:
            raise FileNotFoundError(
                f"{subcollection_name} subcollection not found. "
                f"Run --init-collection --project \"{self.project_name}\" first."
            )

        # Get all notes in subcollection
        notes = self.get_collection_notes(subcollection_key)

        for note in notes:
            title = self.get_note_title_from_html(note['data']['note'])
            if note_title in title:
                content = self.extract_text_from_note_html(note['data']['note'])

                # Parse key=value pairs
                config = {}
                for line in content.split('\n'):
                    line = line.strip()

                    # Skip comments and empty lines
                    if not line or line.startswith('#'):
                        continue

                    # Parse key=value
                    if '=' in line:
                        key, value = line.split('=', 1)
                        key = key.strip()
                        value = value.strip()

                        # Type conversion
                        if value.lower() in ['true', 'false']:
                            config[key] = value.lower() == 'true'
                        elif value.isdigit():
                            config[key] = int(value)
                        elif '.' in value:
                            try:
                                config[key] = float(value)
                            except ValueError:
                                config[key] = value  # Keep as string
                        else:
                            config[key] = value

                return config

        raise FileNotFoundError(
            f"{note_title} not found in {subcollection_name} subcollection. "
            f"Run --init-collection --project \"{self.project_name}\" first."
        )

    def apply_project_config(self, config: Dict[str, Any]):
        """
        Apply loaded configuration to instance attributes with validation.

        Args:
            config: Dict of configuration key-value pairs
        """
        # Validation rules
        def validate_int_range(value, min_val, max_val, name):
            if not isinstance(value, int):
                if self.verbose:
                    print(f"  ⚠️  Invalid {name}: must be integer, got {type(value).__name__}")
                return False
            if not (min_val <= value <= max_val):
                if self.verbose:
                    print(f"  ⚠️  Invalid {name}: {value} (must be {min_val}-{max_val})")
                return False
            return True

        def validate_float_range(value, min_val, max_val, name):
            if not isinstance(value, (int, float)):
                if self.verbose:
                    print(f"  ⚠️  Invalid {name}: must be number, got {type(value).__name__}")
                return False
            if not (min_val <= value <= max_val):
                if self.verbose:
                    print(f"  ⚠️  Invalid {name}: {value} (must be {min_val}-{max_val})")
                return False
            return True

        # Apply each configuration value
        if 'max_workers' in config:
            if validate_int_range(config['max_workers'], 1, 50, 'max_workers'):
                self.max_workers = config['max_workers']

        if 'rate_limit_delay' in config:
            if validate_float_range(config['rate_limit_delay'], 0.0, 10.0, 'rate_limit_delay'):
                self.rate_limit_delay = config['rate_limit_delay']

        if 'general_summary_char_limit' in config:
            if validate_int_range(config['general_summary_char_limit'], 1000, 1000000, 'general_summary_char_limit'):
                self.general_summary_char_limit = config['general_summary_char_limit']

        if 'targeted_summary_char_limit' in config:
            if validate_int_range(config['targeted_summary_char_limit'], 1000, 1000000, 'targeted_summary_char_limit'):
                self.targeted_summary_char_limit = config['targeted_summary_char_limit']

        if 'relevance_threshold' in config:
            if validate_int_range(config['relevance_threshold'], 0, 10, 'relevance_threshold'):
                self.relevance_threshold = config['relevance_threshold']

        if 'max_sources' in config:
            if validate_int_range(config['max_sources'], 1, 1000, 'max_sources'):
                self.max_sources = config['max_sources']

        if 'use_sonnet' in config:
            if isinstance(config['use_sonnet'], bool):
                self.use_sonnet = config['use_sonnet']
                # Update summary model based on use_sonnet
                self.summary_model = self.sonnet_model if self.use_sonnet else self.haiku_model
            elif self.verbose:
                print(f"  ⚠️  Invalid use_sonnet: must be true/false, got {config['use_sonnet']}")

        if 'haiku_model' in config:
            if isinstance(config['haiku_model'], str) and config['haiku_model'].startswith('claude-'):
                self.haiku_model = config['haiku_model']
            elif self.verbose:
                print(f"  ⚠️  Invalid haiku_model: must start with 'claude-'")

        if 'sonnet_model' in config:
            if isinstance(config['sonnet_model'], str) and config['sonnet_model'].startswith('claude-'):
                self.sonnet_model = config['sonnet_model']
            elif self.verbose:
                print(f"  ⚠️  Invalid sonnet_model: must start with 'claude-'")

        if 'gemini_file_search_model' in config:
            if isinstance(config['gemini_file_search_model'], str) and config['gemini_file_search_model'].startswith('gemini-'):
                self.gemini_file_search_model = config['gemini_file_search_model']
            elif self.verbose:
                print(f"  ⚠️  Invalid gemini_file_search_model: must start with 'gemini-'")

        if 'generate_synthesis' in config:
            if isinstance(config['generate_synthesis'], bool):
                self.synthesis_enabled = config['generate_synthesis']
            elif self.verbose:
                print(f"  ⚠️  Invalid generate_synthesis: must be true/false, got {config['generate_synthesis']}")

    def get_note_from_subcollection(
        self,
        collection_key: str,
        note_title: str
    ) -> Optional[Dict]:
        """
        Find and return a note by title from the project subcollection.

        Args:
            collection_key: Parent collection key
            note_title: The exact note title to search for (use _get_*_note_title() methods)

        Returns:
            The note item dict, or None if not found

        Raises:
            FileNotFoundError: If project subcollection doesn't exist
        """
        subcollection_name = self._get_subcollection_name()

        # Get project-specific subcollection
        subcollection_key = self.get_subcollection(collection_key, subcollection_name)
        if not subcollection_key:
            raise FileNotFoundError(
                f"{subcollection_name} subcollection not found. "
                f"Run --init-collection --project \"{self.project_name}\" first."
            )

        # Get all notes in subcollection
        notes = self.get_collection_notes(subcollection_key)

        # Find note by title
        for note in notes:
            title = self.get_note_title_from_html(note['data']['note'])
            if note_title in title:
                return note

        return None

    def _get_items_from_cache(
        self,
        collection_key: str,
        subcollections: Optional[str] = None,
        include_main: bool = False
    ) -> List[Dict]:
        """
        Get items from cache with optional subcollection filtering.

        Args:
            collection_key: Main collection key
            subcollections: Comma-separated subcollection names, "all", or None
            include_main: Include items from main collection (only relevant when subcollections is set)

        Returns:
            List of filtered items from cache
        """
        # If no subcollection filtering, get all items from collection
        if not subcollections:
            cached_items = self.cache_manager.get_cached_items(collection_key)
            # Convert cache format to API format (add 'data' wrapper for metadata)
            return cached_items

        # Handle subcollection filtering
        # For now, pass subcollection names as a list to cache manager
        if subcollections.lower() == "all":
            # Get all subcollections (cache manager will handle excluding ZResearcher subcollection)
            # TODO: Implement in cache manager
            cached_items = self.cache_manager.get_cached_items(collection_key)
        else:
            # Parse comma-separated list
            requested_names = [name.strip() for name in subcollections.split(',')]
            cached_items = self.cache_manager.get_cached_items(collection_key, subcollections=requested_names)

        return cached_items

    def has_note_with_prefix(self, item_key: str, prefix: str) -> bool:
        """
        Check if an item already has a note starting with a specific prefix.
        Cache-aware version that delegates to cache when available.

        Args:
            item_key: The key of the parent item
            prefix: The prefix to search for (e.g., "【ZResearcher Summary:")

        Returns:
            True if the item has a note with that prefix
        """
        # Use cache if available
        if self.use_cache and self.cache_manager:
            notes = self.cache_manager.get_cached_notes_for_item(item_key, title_prefix=prefix)
            return len(notes) > 0

        # Otherwise use API
        return super().has_note_with_prefix(item_key, prefix)

    def get_collection_notes(self, collection_key: str) -> List[Dict]:
        """
        Get all standalone notes in a collection.
        Cache-aware version that delegates to cache when available.

        Args:
            collection_key: Collection key

        Returns:
            List of note items
        """
        # Use cache if available
        if self.use_cache and self.cache_manager:
            notes = self.cache_manager.get_standalone_notes(collection_key)
            # Convert cache format to API format if needed
            return notes

        # Otherwise use API
        return super().get_collection_notes(collection_key)

    def get_note_with_prefix(self, item_key: str, prefix: str) -> Optional[str]:
        """
        Get the content of a note starting with a specific prefix.
        Cache-aware version that delegates to cache when available.

        Args:
            item_key: The key of the parent item
            prefix: The prefix to search for (e.g., "【ZResearcher Summary:")

        Returns:
            The HTML content of the note, or None if not found
        """
        # Use cache if available
        if self.use_cache and self.cache_manager:
            notes = self.cache_manager.get_cached_notes_for_item(item_key, title_prefix=prefix)
            if notes:
                # Return the HTML content from the first matching note
                # Note: Cache stores HTML in the 'content' field
                return notes[0].get('content', notes[0].get('note', ''))

        # Otherwise use API
        return super().get_note_with_prefix(item_key, prefix)

    def get_subcollection(self, parent_collection_key: str, subcollection_name: str) -> Optional[str]:
        """
        Get a subcollection key by name within a parent collection.
        Cache-aware version that delegates to cache when available.

        Args:
            parent_collection_key: Key of parent collection
            subcollection_name: Name of subcollection to find

        Returns:
            Subcollection key or None if not found
        """
        # Use cache if available
        if self.use_cache and self.cache_manager:
            try:
                conn = self.cache_manager.get_connection()
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT collection_key FROM collections
                    WHERE parent_key = ? AND name = ?
                """, (parent_collection_key, subcollection_name))
                result = cursor.fetchone()
                conn.close()

                if result:
                    return result['collection_key']
                return None
            except Exception as e:
                if self.verbose:
                    print(f"  ⚠️  Cache error, falling back to API: {e}")
                # Fall through to API

        # Otherwise use API
        return super().get_subcollection(parent_collection_key, subcollection_name)

    def get_items_to_process(
        self,
        collection_key: str,
        subcollections: Optional[str] = None,
        include_main: bool = False
    ) -> List[Dict]:
        """
        Get items to process with optional subcollection filtering.

        Args:
            collection_key: Main collection key
            subcollections: Comma-separated subcollection names, "all", or None
            include_main: Include items from main collection (only relevant when subcollections is set)

        Returns:
            List of filtered items to process

        Raises:
            ValueError: If specified subcollection names not found
        """
        # Use cache if available
        if self.use_cache and self.cache_manager:
            return self._get_items_from_cache(collection_key, subcollections, include_main)

        # If no subcollection filtering, return all items from collection
        if not subcollections:
            return self.get_collection_items(collection_key)

        # Get all subcollections of the parent collection
        try:
            all_subcollections = self.zot.collections_sub(collection_key)
        except Exception as e:
            print(f"  ❌ Error fetching subcollections: {e}")
            return []

        # Build map of subcollection names to keys
        subcollection_map = {}
        for subcoll in all_subcollections:
            subcoll_name = subcoll['data']['name']
            subcoll_key = subcoll['key']
            subcollection_map[subcoll_name] = subcoll_key

        # Exclude the ZResearcher project subcollection if it exists
        if self.project_name:
            zresearcher_subcoll_name = self._get_subcollection_name()
            if zresearcher_subcoll_name in subcollection_map:
                del subcollection_map[zresearcher_subcoll_name]

        # Determine which subcollections to include
        target_subcollection_keys = set()

        if subcollections.lower() == "all":
            # Include all subcollections (except ZResearcher project subcollection)
            target_subcollection_keys = set(subcollection_map.values())
            if self.verbose:
                print(f"  📁 Filtering to all subcollections ({len(target_subcollection_keys)} total)")
        else:
            # Parse comma-separated list
            requested_names = [name.strip() for name in subcollections.split(',')]
            for name in requested_names:
                if name not in subcollection_map:
                    available = ', '.join(f'"{n}"' for n in sorted(subcollection_map.keys()))
                    raise ValueError(
                        f"Subcollection '{name}' not found. "
                        f"Available subcollections: {available}"
                    )
                target_subcollection_keys.add(subcollection_map[name])

            if self.verbose:
                print(f"  📁 Filtering to subcollections: {', '.join(requested_names)}")

        # Get items from target subcollections
        filtered_items = []
        seen_keys = set()  # Track items to avoid duplicates

        # Get items from each target subcollection
        for subcoll_key in target_subcollection_keys:
            try:
                # Use everything() to handle pagination and fetch all items (no 100-item limit)
                subcoll_items = self.zot.everything(self.zot.collection_items_top(subcoll_key))
                for item in subcoll_items:
                    item_key = item['key']
                    if item_key not in seen_keys:
                        filtered_items.append(item)
                        seen_keys.add(item_key)
            except Exception as e:
                print(f"  ⚠️  Error fetching items from subcollection: {e}")

        # If include_main=True, also get items from main collection that are NOT in any subcollection
        if include_main:
            try:
                main_items = self.get_collection_items(collection_key)
                for item in main_items:
                    item_key = item['key']
                    item_collections = item['data'].get('collections', [])

                    # Check if item is in main collection only (not in any subcollection)
                    in_any_subcollection = any(
                        coll_key in subcollection_map.values()
                        for coll_key in item_collections
                    )

                    # Add if not already seen and not in any subcollection
                    if item_key not in seen_keys and not in_any_subcollection:
                        filtered_items.append(item)
                        seen_keys.add(item_key)
            except Exception as e:
                print(f"  ⚠️  Error fetching items from main collection: {e}")

        if self.verbose:
            print(f"  ✅ Found {len(filtered_items)} items after subcollection filtering")

        return filtered_items

    def load_note_from_subcollection(
        self,
        collection_key: str,
        note_title: str,
        check_todo: bool = True,
        remove_title_line: bool = True,
        remove_footer: bool = True,
        operation_name: str = None
    ) -> str:
        """
        Load and clean text content from a note in the project subcollection.

        This is a generic method that consolidates common note loading logic
        used across multiple workflows (build, query, file search).

        Args:
            collection_key: Parent collection key
            note_title: The exact note title to search for (use _get_*_note_title() methods)
            check_todo: If True, raise error if note contains [TODO: markers
            remove_title_line: If True, remove first line if it matches note_title
            remove_footer: If True, remove content after '---' separator
            operation_name: Human-readable operation name for error messages (e.g., "building summaries")

        Returns:
            Cleaned note content as text

        Raises:
            FileNotFoundError: If subcollection or note not found
            ValueError: If note contains [TODO: markers and check_todo=True
        """
        subcollection_name = self._get_subcollection_name()

        # Find the note
        note = self.get_note_from_subcollection(collection_key, note_title)
        if not note:
            error_msg = f"{note_title} not found in {subcollection_name} subcollection. "
            error_msg += f"Run --init-collection --project \"{self.project_name}\" first."
            if operation_name:
                error_msg += f" Edit the note before {operation_name}."
            raise FileNotFoundError(error_msg)

        # Extract text content
        content = self.extract_text_from_note_html(note['data']['note'])

        # Check for template placeholder
        if check_todo and '[TODO:' in content:
            error_msg = f"{note_title} still contains template. "
            error_msg += "Please edit the note in Zotero"
            if operation_name:
                error_msg += f" before {operation_name}"
            error_msg += "."
            raise ValueError(error_msg)

        # Remove footer separator if present
        if remove_footer and '---' in content:
            content = content.split('---')[0]

        # Remove title line if present
        if remove_title_line:
            lines = content.split('\n')
            if lines and note_title in lines[0]:
                content = '\n'.join(lines[1:])

        return content.strip()

    def update_note_in_subcollection(
        self,
        collection_key: str,
        note_title: str,
        new_content: str,
        preserve_formatting: bool = True
    ) -> None:
        """
        Update a note's content in the project subcollection.

        Args:
            collection_key: Parent collection key
            note_title: The exact note title to search for
            new_content: New text content (markdown)
            preserve_formatting: If True, wrap in code block to preserve formatting

        Raises:
            FileNotFoundError: If subcollection or note not found
        """
        subcollection_name = self._get_subcollection_name()

        # Find the note
        note = self.get_note_from_subcollection(collection_key, note_title)
        if not note:
            raise FileNotFoundError(
                f"{note_title} not found in {subcollection_name} subcollection."
            )

        # Convert to HTML
        if preserve_formatting:
            # Wrap in code block to preserve formatting
            updated_html = self.markdown_to_html(f"```\n{new_content}\n```")
        else:
            updated_html = self.markdown_to_html(new_content)

        # Update the note
        note['data']['note'] = updated_html
        self.zot.update_item(note)

        if self.verbose:
            print(f"  ✅ Updated {note_title}")
